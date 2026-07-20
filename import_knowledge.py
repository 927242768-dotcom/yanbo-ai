"""将 PDF、Word、图片、文本和代码资料导入彦博本地知识库。"""

from __future__ import annotations

import argparse
import hashlib
import re
import sys
from pathlib import Path
from typing import Iterable

from console_utils import configure_utf8_console
from image_understanding import ImageTextRecognizer


ROOT = Path(__file__).resolve().parent
KNOWLEDGE_DIR = ROOT / "knowledge"
GENERATED_DIR = KNOWLEDGE_DIR / "generated"
TEXT_SUFFIXES = {
    ".txt",
    ".md",
    ".csv",
    ".json",
    ".jsonl",
    ".py",
    ".c",
    ".h",
    ".cpp",
    ".hpp",
    ".js",
    ".ts",
    ".java",
    ".sql",
}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".bmp"}
SUPPORTED_SUFFIXES = TEXT_SUFFIXES | IMAGE_SUFFIXES | {".pdf", ".docx"}
MAX_SOURCE_BYTES = 100 * 1024 * 1024


def _safe_name(value: str) -> str:
    name = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff._-]+", "-", value).strip("-._")
    return name[:80] or "资料"


def _source_files(paths: Iterable[str]) -> list[Path]:
    files: list[Path] = []
    seen: set[Path] = set()
    for raw in paths:
        path = Path(raw).expanduser().resolve()
        candidates = path.rglob("*") if path.is_dir() else [path]
        for candidate in candidates:
            if not candidate.is_file() or candidate.suffix.lower() not in SUPPORTED_SUFFIXES:
                continue
            resolved = candidate.resolve()
            if resolved in seen:
                continue
            seen.add(resolved)
            files.append(resolved)
    return sorted(files)


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("缺少PDF解析组件，请先运行 00_setup.bat。") from exc
    reader = PdfReader(str(path))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            text = (page.extract_text() or "").strip()
        except Exception as exc:  # 单页异常不应导致整份资料丢失。
            text = f"[第{index}页解析失败：{exc}]"
        if text:
            pages.append(f"## 第{index}页\n\n{text}")
    return "\n\n".join(pages)


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少Word解析组件，请先运行 00_setup.bat。") from exc
    document = Document(str(path))
    rows = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table_index, table in enumerate(document.tables, start=1):
        table_rows = []
        for row in table.rows:
            cells = [re.sub(r"\s+", " ", cell.text).strip() for cell in row.cells]
            if any(cells):
                table_rows.append(" | ".join(cells))
        if table_rows:
            rows.append(f"\n[表格{table_index}]\n" + "\n".join(table_rows))
    return "\n\n".join(rows)


def _read_image(path: Path, recognizer: ImageTextRecognizer) -> str:
    result = recognizer.recognize_bytes(path.read_bytes())
    return (
        f"图片文字识别置信度：{result.confidence:.1%}\n"
        f"图片尺寸：{result.width}×{result.height}\n\n"
        f"{result.text}"
    )


def _extract(path: Path, recognizer: ImageTextRecognizer | None) -> str:
    suffix = path.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return _read_text(path)
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    if suffix in IMAGE_SUFFIXES:
        if recognizer is None:
            recognizer = ImageTextRecognizer()
        return _read_image(path, recognizer)
    raise ValueError(f"不支持的文件格式：{suffix}")


def _output_path(path: Path, digest: str) -> Path:
    name = _safe_name(path.stem)
    return GENERATED_DIR / f"{name}-{digest[:12]}.md"


def import_file(path: Path, recognizer: ImageTextRecognizer | None = None) -> tuple[Path, bool]:
    size = path.stat().st_size
    if size <= 0:
        raise ValueError("文件为空")
    if size > MAX_SOURCE_BYTES:
        raise ValueError("文件超过100MB上限")
    raw = path.read_bytes()
    digest = hashlib.sha256(raw).hexdigest()
    output = _output_path(path, digest)
    if output.exists():
        return output, False
    text = _extract(path, recognizer).strip()
    if not text:
        raise ValueError("没有提取到可检索文字")
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    content = (
        f"# {path.name}\n\n"
        f"- 原始路径：`{path}`\n"
        f"- 文件哈希：`{digest}`\n\n"
        f"---\n\n{text}\n"
    )
    output.write_text(content, encoding="utf-8")
    return output, True


def main() -> None:
    configure_utf8_console()
    parser = argparse.ArgumentParser(
        description="把资料转换为可被彦博自动检索的本地知识库文本"
    )
    parser.add_argument("paths", nargs="+", help="文件或文件夹路径，可一次提供多个")
    args = parser.parse_args()

    files = _source_files(args.paths)
    if not files:
        print("没有找到支持的资料文件。")
        raise SystemExit(1)

    recognizer = ImageTextRecognizer() if any(path.suffix.lower() in IMAGE_SUFFIXES for path in files) else None
    succeeded = 0
    skipped = 0
    failed = 0
    for index, path in enumerate(files, start=1):
        try:
            output, created = import_file(path, recognizer)
            if created:
                succeeded += 1
                print(f"[{index}/{len(files)}] 已导入：{path.name} -> {output.relative_to(ROOT)}")
            else:
                skipped += 1
                print(f"[{index}/{len(files)}] 已存在，跳过：{path.name}")
        except Exception as exc:
            failed += 1
            print(f"[{index}/{len(files)}] 导入失败：{path}：{exc}", file=sys.stderr)

    print(f"\n导入完成：新增{succeeded}，跳过{skipped}，失败{failed}。")
    print("新资料会在下一次提问时自动进入检索，无需重启或重新训练。")
    if failed and not succeeded:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
